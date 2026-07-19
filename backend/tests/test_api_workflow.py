from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient


def make_resume() -> BytesIO:
    document = Document()
    document.add_heading("Taylor Candidate", level=1)
    document.add_paragraph(
        "Backend engineer with Python, FastAPI, React, PostgreSQL, Docker, RAG, and ChromaDB experience."
    )
    content = BytesIO()
    document.save(content)
    content.seek(0)
    return content


def register(client: TestClient, email: str) -> dict[str, str]:
    response = client.post(
        "/auth/register",
        json={"email": email, "password": "password123", "name": "Test User"},
    )
    assert response.status_code == 201, response.text
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def test_health(client: TestClient) -> None:
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_complete_application_workflow_requires_user_confirmation(client: TestClient) -> None:
    headers = register(client, "taylor@example.com")
    upload = client.post(
        "/resumes/upload",
        headers=headers,
        files={
            "file": (
                "taylor-resume.docx",
                make_resume(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 200, upload.text
    resume = upload.json()
    assert resume["file_name"] == "taylor-resume.docx"

    imported = client.post(
        "/jobs/import-text",
        headers=headers,
        json={
            "company": "Example Labs",
            "title": "Backend Engineer",
            "location": "London",
            "url": "https://example.com/jobs/backend-engineer",
            "raw_description": (
                "We require Python, FastAPI, React, PostgreSQL, Docker, RAG, and ChromaDB skills. "
                "Build reliable APIs and collaborate across the full product lifecycle."
            ),
        },
    )
    assert imported.status_code == 200, imported.text
    job = imported.json()

    matched = client.post(f"/jobs/{job['id']}/match/{resume['id']}", headers=headers)
    assert matched.status_code == 200, matched.text
    match_payload = matched.json()
    assert match_payload["match"]["score"] >= 70
    application_id = match_payload["application_id"]

    pending = client.get("/applications/pending", headers=headers)
    assert pending.status_code == 200
    assert [item["id"] for item in pending.json()] == [application_id]
    assert pending.json()[0]["confirmed_by_user"] is False

    detail = client.get(f"/applications/{application_id}", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["job"]["company"] == "Example Labs"
    assert detail.json()["materials"][0]["material_type"] == "cover_letter"

    material = detail.json()["materials"][0]
    updated = client.put(
        f"/materials/{material['id']}",
        headers=headers,
        json={"content": "Reviewed and tailored cover letter."},
    )
    assert updated.status_code == 200
    assert updated.json()["content"] == "Reviewed and tailored cover letter."

    confirmed = client.post(f"/applications/{application_id}/confirm", headers=headers)
    assert confirmed.status_code == 200
    assert confirmed.json()["status"] == "applied"
    assert confirmed.json()["confirmed_by_user"] is True

    assert client.get("/applications/pending", headers=headers).json() == []
    duplicate_confirmation = client.post(f"/applications/{application_id}/confirm", headers=headers)
    assert duplicate_confirmation.status_code == 409


def test_match_rejects_unknown_resources(client: TestClient) -> None:
    headers = register(client, "unknown@example.com")
    response = client.post(
        "/jobs/00000000-0000-0000-0000-000000000000/match/00000000-0000-0000-0000-000000000000",
        headers=headers,
    )
    assert response.status_code == 404
    assert response.json()["detail"] == "Job or resume not found."


def test_protected_endpoints_require_authentication(client: TestClient) -> None:
    assert client.get("/resumes").status_code == 401
    assert client.get("/jobs").status_code == 401
    assert client.get("/applications/pending").status_code == 401
    assert client.get("/dashboard/summary").status_code == 401


def test_login_and_current_user(client: TestClient) -> None:
    headers = register(client, "login@example.com")
    me = client.get("/auth/me", headers=headers)
    assert me.status_code == 200
    assert me.json()["email"] == "login@example.com"
    assert "password_hash" not in me.json()

    invalid = client.post("/auth/login", json={"email": "login@example.com", "password": "wrong-password"})
    assert invalid.status_code == 401

    login = client.post("/auth/login", json={"email": "login@example.com", "password": "password123"})
    assert login.status_code == 200
    assert login.json()["token_type"] == "bearer"
    assert login.json()["access_token"]


def test_users_cannot_access_each_others_data(client: TestClient) -> None:
    owner_headers = register(client, "owner@example.com")
    other_headers = register(client, "other@example.com")

    upload = client.post(
        "/resumes/upload",
        headers=owner_headers,
        files={
            "file": (
                "owner.docx",
                make_resume(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    resume_id = upload.json()["id"]
    imported = client.post(
        "/jobs/import-text",
        headers=owner_headers,
        json={
            "company": "Private Company",
            "title": "Private Role",
            "raw_description": "Python FastAPI PostgreSQL Docker backend engineering role.",
        },
    )
    job_id = imported.json()["id"]

    assert client.get("/resumes", headers=other_headers).json() == []
    assert client.get("/jobs", headers=other_headers).json() == []
    assert client.get("/dashboard/summary", headers=owner_headers).json()["total_jobs"] == 1
    assert client.get("/dashboard/summary", headers=other_headers).json()["total_jobs"] == 0
    cross_user_match = client.post(f"/jobs/{job_id}/match/{resume_id}", headers=other_headers)
    assert cross_user_match.status_code == 404

    owner_match = client.post(f"/jobs/{job_id}/match/{resume_id}", headers=owner_headers)
    assert owner_match.status_code == 200
    application_id = owner_match.json()["application_id"]
    material_id = owner_match.json()["cover_letter_id"]

    assert client.get(f"/applications/{application_id}", headers=other_headers).status_code == 404
    assert client.get(f"/materials/{material_id}", headers=other_headers).status_code == 404
    assert client.get("/materials", headers=other_headers).json() == []
    assert client.post(f"/applications/{application_id}/confirm", headers=other_headers).status_code == 404
